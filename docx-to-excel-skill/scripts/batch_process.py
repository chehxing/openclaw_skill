#!/usr/bin/env python3
"""
批量处理多个Word文档并提取数据到Excel
支持合并多个文档数据到单个Excel文件
"""

import argparse
import sys
from pathlib import Path
from typing import List, Dict, Any
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment


def find_docx_files(input_dir: str, pattern: str = "*.docx") -> List[Path]:
    """查找指定目录下的所有.docx文件"""
    input_path = Path(input_dir)
    if not input_path.exists():
        print(f"错误: 输入目录不存在: {input_dir}")
        return []

    docx_files = list(input_path.glob(pattern))
    docx_files.extend(input_path.glob("*.DOCX"))  # 大写扩展名

    # 按文件名排序
    docx_files.sort()
    return docx_files


def extract_document_info(doc_path: Path) -> Dict[str, Any]:
    """提取文档基本信息"""
    try:
        doc = Document(str(doc_path))
        info = {
            '文件名': doc_path.name,
            '文件路径': str(doc_path),
            '文件大小': f"{doc_path.stat().st_size / 1024:.1f} KB",
            '段落数': len(doc.paragraphs),
            '表格数': len(doc.tables),
            '图片数': len(doc.inline_shapes),
            '页数': '未知'  # python-docx不直接支持页数
        }

        # 尝试提取文档属性
        if doc.core_properties:
            if doc.core_properties.title:
                info['标题'] = doc.core_properties.title
            if doc.core_properties.author:
                info['作者'] = doc.core_properties.author
            if doc.core_properties.created:
                info['创建日期'] = doc.core_properties.created.strftime('%Y-%m-%d')
            if doc.core_properties.modified:
                info['修改日期'] = doc.core_properties.modified.strftime('%Y-%m-%d')

        # 提取前几个段落作为摘要
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            info['首段'] = paragraphs[0][:100] + "..." if len(paragraphs[0]) > 100 else paragraphs[0]

        return info

    except Exception as e:
        print(f"处理文件失败 {doc_path.name}: {e}")
        return {'文件名': doc_path.name, '错误': str(e)}


def extract_tables_summary(doc_path: Path) -> List[Dict[str, Any]]:
    """提取文档中的表格摘要信息"""
    try:
        doc = Document(str(doc_path))
        tables_summary = []

        for table_idx, table in enumerate(doc.tables):
            table_info = {
                '文档': doc_path.name,
                '表格索引': table_idx + 1,
                '行数': len(table.rows),
                '列数': len(table.columns) if table.columns else 0,
                '总单元格数': len(table.rows) * (len(table.columns) if table.columns else 0)
            }

            # 提取表头（第一行）
            if table.rows:
                headers = []
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        headers.append(header_text)
                table_info['表头'] = ', '.join(headers) if headers else '无'

            tables_summary.append(table_info)

        return tables_summary

    except Exception as e:
        print(f"提取表格摘要失败 {doc_path.name}: {e}")
        return []


def create_summary_excel(documents_info: List[Dict[str, Any]],
                        tables_summary: List[Dict[str, Any]],
                        output_path: str) -> bool:
    """创建汇总Excel文件"""
    try:
        wb = Workbook()

        # 文档信息工作表
        if documents_info:
            ws_docs = wb.active
            ws_docs.title = "文档信息"

            # 获取所有字段
            all_fields = set()
            for info in documents_info:
                all_fields.update(info.keys())

            fields = list(all_fields)
            ws_docs.append(fields)

            # 写入数据
            for info in documents_info:
                row = [info.get(field, '') for field in fields]
                ws_docs.append(row)

            # 设置列宽和格式
            for col_idx, field in enumerate(fields, 1):
                ws_docs.column_dimensions[chr(64 + col_idx)].width = max(len(field), 15)

            # 表头加粗
            for cell in ws_docs[1]:
                cell.font = Font(bold=True)

        # 表格摘要工作表
        if tables_summary:
            ws_tables = wb.create_sheet(title="表格摘要")

            if tables_summary:
                fields = list(tables_summary[0].keys())
                ws_tables.append(fields)

                for table_info in tables_summary:
                    row = [table_info.get(field, '') for field in fields]
                    ws_tables.append(row)

                # 设置列宽
                for col_idx, field in enumerate(fields, 1):
                    ws_tables.column_dimensions[chr(64 + col_idx)].width = max(len(field), 15)

                # 表头加粗
                for cell in ws_tables[1]:
                    cell.font = Font(bold=True)

        # 统计信息工作表
        ws_stats = wb.create_sheet(title="统计信息")
        stats_data = [
            ["统计项", "数值"],
            ["处理文档数", len(documents_info)],
            ["发现表格总数", len(tables_summary)],
            ["成功处理文档", sum(1 for d in documents_info if '错误' not in d)],
            ["失败文档", sum(1 for d in documents_info if '错误' in d)]
        ]

        for row in stats_data:
            ws_stats.append(row)

        # 设置格式
        for col in ['A', 'B']:
            ws_stats.column_dimensions[col].width = 20

        for cell in ws_stats[1]:
            cell.font = Font(bold=True)

        # 保存文件
        wb.save(output_path)
        return True

    except Exception as e:
        print(f"创建Excel文件失败: {e}")
        return False


def create_individual_excels(documents_info: List[Dict[str, Any]],
                           output_dir: str) -> bool:
    """为每个文档创建单独的Excel文件"""
    try:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        for info in documents_info:
            if '错误' in info:
                continue

            filename = info.get('文件名', 'unknown').replace('.docx', '.xlsx')
            filepath = output_path / filename

            wb = Workbook()
            ws = wb.active
            ws.title = "文档信息"

            # 写入文档信息
            ws.append(["属性", "值"])
            for key, value in info.items():
                if key != '文件名':  # 文件名已经在文件名中体现了
                    ws.append([key, str(value)])

            # 设置格式
            ws.column_dimensions['A'].width = 20
            ws.column_dimensions['B'].width = 40

            for cell in ws[1]:
                cell.font = Font(bold=True)

            wb.save(str(filepath))

        return True

    except Exception as e:
        print(f"创建单独Excel文件失败: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(description='批量处理Word文档并提取信息到Excel')
    parser.add_argument('--input-dir', '-i', required=True, help='输入目录路径')
    parser.add_argument('--output', '-o', default='batch_output.xlsx', help='输出Excel文件路径')
    parser.add_argument('--merge', '-m', action='store_true', default=True,
                       help='合并所有文档数据到单个文件 (默认: True)')
    parser.add_argument('--no-merge', action='store_false', dest='merge',
                       help='为每个文档创建单独文件')
    parser.add_argument('--pattern', '-p', default='*.docx', help='文件匹配模式 (默认: *.docx)')
    parser.add_argument('--verbose', '-v', action='store_true', help='显示详细信息')

    args = parser.parse_args()

    # 查找文档文件
    if args.verbose:
        print(f"在目录中查找文件: {args.input_dir}")
        print(f"使用模式: {args.pattern}")

    docx_files = find_docx_files(args.input_dir, args.pattern)

    if not docx_files:
        print(f"错误: 未找到匹配的.docx文件: {args.input_dir}/{args.pattern}")
        sys.exit(1)

    if args.verbose:
        print(f"找到 {len(docx_files)} 个文档文件")

    # 处理每个文档
    documents_info = []
    all_tables_summary = []

    for doc_path in docx_files:
        if args.verbose:
            print(f"处理: {doc_path.name}")

        # 提取文档信息
        doc_info = extract_document_info(doc_path)
        documents_info.append(doc_info)

        # 提取表格摘要
        tables_info = extract_tables_summary(doc_path)
        all_tables_summary.extend(tables_info)

        if args.verbose and '错误' in doc_info:
            print(f"  错误: {doc_info['错误']}")

    # 创建输出
    if args.merge:
        # 合并到单个文件
        if args.verbose:
            print(f"创建合并文件: {args.output}")

        success = create_summary_excel(documents_info, all_tables_summary, args.output)

        if success:
            print(f"批量处理完成! 结果保存到: {args.output}")
            print(f"处理了 {len(docx_files)} 个文档，找到 {len(all_tables_summary)} 个表格")
        else:
            print("创建合并文件失败!")
            sys.exit(1)
    else:
        # 创建单独文件
        output_dir = Path(args.output).parent / "individual_excels"
        if args.verbose:
            print(f"创建单独文件到目录: {output_dir}")

        success = create_individual_excels(documents_info, str(output_dir))

        if success:
            print(f"批量处理完成! 结果保存到目录: {output_dir}")
            print(f"为 {len([d for d in documents_info if '错误' not in d])} 个文档创建了单独文件")
        else:
            print("创建单独文件失败!")
            sys.exit(1)

    sys.exit(0)


if __name__ == "__main__":
    main()