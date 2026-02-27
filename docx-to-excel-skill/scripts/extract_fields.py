#!/usr/bin/env python3
"""
根据自定义字段映射从Word文档提取数据到Excel
支持多种提取类型：段落、表格、正则表达式
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document
from openpyxl import Workbook, load_workbook


class FieldExtractor:
    """字段提取器基类"""

    def __init__(self, doc: Document):
        self.doc = doc

    def extract(self, field_config: Dict[str, Any]) -> str:
        """提取字段值"""
        raise NotImplementedError


class ParagraphExtractor(FieldExtractor):
    """段落提取器"""

    def extract(self, field_config: Dict[str, Any]) -> str:
        location = field_config.get('location', '')
        if not location:
            return ""

        try:
            # 解析位置格式：标题级别[索引]
            match = re.match(r'标题(\d+)\[(\d+)\]', location)
            if match:
                level = int(match.group(1))
                index = int(match.group(2))

                # 查找指定标题级别的段落
                found_count = 0
                for para in self.doc.paragraphs:
                    if para.style.name.startswith(f'Heading {level}'):
                        if found_count == index:
                            return para.text.strip()
                        found_count += 1

            # 如果未找到，尝试其他定位方式
            return self._find_by_text_pattern(field_config)
        except Exception as e:
            print(f"段落提取错误: {e}")
            return field_config.get('default', '')

    def _find_by_text_pattern(self, field_config: Dict[str, Any]) -> str:
        """通过文本模式查找"""
        pattern = field_config.get('pattern')
        if not pattern:
            return ""

        for para in self.doc.paragraphs:
            if re.search(pattern, para.text):
                # 提取匹配的部分
                match = re.search(pattern, para.text)
                if match and match.groups():
                    return match.group(1).strip()
                return para.text.strip()
        return field_config.get('default', '')


class TableExtractor(FieldExtractor):
    """表格提取器"""

    def extract(self, field_config: Dict[str, Any]) -> str:
        table_index = field_config.get('table_index', 0)
        row = field_config.get('row', 0)
        column = field_config.get('column', 0)

        try:
            if table_index >= len(self.doc.tables):
                return field_config.get('default', '')

            table = self.doc.tables[table_index]
            if row >= len(table.rows) or column >= len(table.rows[row].cells):
                return field_config.get('default', '')

            cell = table.rows[row].cells[column]
            return cell.text.strip()
        except Exception as e:
            print(f"表格提取错误: {e}")
            return field_config.get('default', '')


class RegexExtractor(FieldExtractor):
    """正则表达式提取器"""

    def extract(self, field_config: Dict[str, Any]) -> str:
        pattern = field_config.get('pattern')
        if not pattern:
            return field_config.get('default', '')

        try:
            # 在整个文档中搜索
            full_text = '\n'.join([para.text for para in self.doc.paragraphs])
            match = re.search(pattern, full_text, re.IGNORECASE | re.MULTILINE)

            if match:
                # 返回第一个捕获组，如果没有捕获组则返回整个匹配
                if match.groups():
                    return match.group(1).strip()
                return match.group(0).strip()
            return field_config.get('default', '')
        except re.error as e:
            print(f"正则表达式错误: {e}")
            return field_config.get('default', '')


def load_field_config(config_path: Optional[str] = None,
                     fields_str: Optional[str] = None) -> List[Dict[str, Any]]:
    """加载字段配置"""
    if config_path:
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get('fields', [])
        except Exception as e:
            print(f"加载配置文件失败: {e}")
            return []

    elif fields_str:
        # 从命令行参数创建简单配置
        fields = []
        for field_name in fields_str.split(','):
            fields.append({
                'name': field_name.strip(),
                'type': 'paragraph',
                'location': '',
                'default': ''
            })
        return fields

    return []


def extract_fields(doc_path: str, field_configs: List[Dict[str, Any]]) -> Dict[str, str]:
    """提取所有字段"""
    try:
        doc = Document(doc_path)
        results = {}

        # 创建提取器映射
        extractors = {
            'paragraph': ParagraphExtractor(doc),
            'table': TableExtractor(doc),
            'regex': RegexExtractor(doc)
        }

        for config in field_configs:
            field_name = config.get('name', '')
            field_type = config.get('type', 'paragraph')

            if not field_name:
                continue

            extractor = extractors.get(field_type)
            if extractor:
                value = extractor.extract(config)
                results[field_name] = value
            else:
                print(f"警告: 未知的字段类型: {field_type}")
                results[field_name] = config.get('default', '')

        return results

    except Exception as e:
        print(f"提取字段失败: {e}")
        return {}


def save_to_excel(data: Dict[str, str], output_path: str,
                 template_path: Optional[str] = None) -> bool:
    """保存数据到Excel"""
    try:
        if template_path and Path(template_path).exists():
            # 使用模板
            wb = load_workbook(template_path)
            ws = wb.active

            # 查找并填充字段占位符
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for field_name, field_value in data.items():
                            placeholder = f"{{{field_name}}}"
                            if placeholder in cell.value:
                                cell.value = cell.value.replace(placeholder, field_value)
        else:
            # 创建新工作簿
            wb = Workbook()
            ws = wb.active
            ws.title = "Extracted Data"

            # 写入数据
            ws.append(["字段名称", "字段值"])
            for field_name, field_value in data.items():
                ws.append([field_name, field_value])

            # 设置列宽
            ws.column_dimensions['A'].width = 20
            ws.column_dimensions['B'].width = 40

        # 保存文件
        wb.save(output_path)
        return True

    except Exception as e:
        print(f"保存Excel文件失败: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(description='从Word文档提取指定字段到Excel')
    parser.add_argument('input', help='输入Word文档路径 (.docx)')
    parser.add_argument('--output', '-o', required=True, help='输出Excel文件路径')
    parser.add_argument('--fields', '-f', help='要提取的字段列表，逗号分隔')
    parser.add_argument('--config', '-c', help='字段映射配置文件路径 (JSON)')
    parser.add_argument('--template', '-t', help='Excel模板文件路径')
    parser.add_argument('--verbose', '-v', action='store_true', help='显示详细信息')

    args = parser.parse_args()

    # 检查输入文件
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误: 输入文件不存在: {args.input}")
        sys.exit(1)

    # 加载字段配置
    if not args.fields and not args.config:
        print("错误: 必须指定 --fields 或 --config 参数")
        sys.exit(1)

    field_configs = load_field_config(args.config, args.fields)

    if not field_configs:
        print("错误: 未加载到有效的字段配置")
        sys.exit(1)

    if args.verbose:
        print(f"加载了 {len(field_configs)} 个字段配置")

    # 提取字段
    if args.verbose:
        print(f"正在提取字段...")

    extracted_data = extract_fields(str(input_path), field_configs)

    if args.verbose:
        print(f"提取结果:")
        for field, value in extracted_data.items():
            print(f"  {field}: {value}")

    # 保存到Excel
    success = save_to_excel(extracted_data, args.output, args.template)

    if success:
        print(f"数据已保存到: {args.output}")
        sys.exit(0)
    else:
        print("保存失败!")
        sys.exit(1)


if __name__ == "__main__":
    main()